from docxPdfImage import *
from docx.enum.text import WD_COLOR_INDEX
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re

"""
Testing iter_block_items()
"""
def iter_block_items(parent):
    """
    Tạo tham chiếu đến từng đoạn và bảng con trong file doc, theo thứ tự tài liệu. 
    Mỗi giá trị trả về là một thể hiện của Bảng hoặc Đoạn văn. 
    'parent' thường là một tham chiếu đến một chính Đối tượng tài liệu, 
    hoạt động cho đối tượng _Cell | đoạn văn | bảng
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def color_string(key,countKey,p1,p):
##    tô vàng key
##    input: key, số thứ tự key, đoạn văn chứa key, đoạn mới chứa key được tô vàng
##    output: đoạn văn đã được tô vàng, số thứ tự key
    substrings = p1.split(key) # split đoạn
    for substring in substrings[:-1]:
        countKey += 1
        p.add_run(substring)
        font = p.add_run(key).font.highlight_color = WD_COLOR_INDEX.YELLOW # tô vàng key
        count = str(countKey)
        font = p.add_run(count).font.highlight_color = WD_COLOR_INDEX.RED # tô đỏ số thứ tự của key
    p.add_run(substrings[-1])
    return countKey
    
def findColor(filename,key,newName):
##    tìm và tô vàng key
##    input: file cần xử ký, key cần tìm và tô vàng
##    output: file đã tô vàng và đánh thứ tự cho key
    countKey = 0 # khởi tạo số thứ tự key
    doc = Document(filename)
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            p1 = block.text
            match = re.search(key,p1,re.IGNORECASE)
            if match: #so khớp không phân biệt hoa thường
                block.text = ""
                countKey=color_string(match.group(),countKey,p1,block)
        else:
            for row in block.rows:
                for p in row.cells:
                    p1 = p.text
                    match = re.search(key,p1,re.IGNORECASE)
                    if match: #so khớp không phân biệt hoa thường
                        p.text = ""
                        p = p.add_paragraph()
                        countKey=color_string(match.group(),countKey,p1,p)
    doc.save(newName)
    return countKey

def replace_string(key,value,numberList,countKey,p):
##    split đoạn văn và key thành list
##    kiểm tra xem key có xuất hiện trong đoạn không
##    đếm số lần xuất hiện của key, nếu thứ tự nằm trong numberList thì thay đổi key
##    thay đổi: xóa các từ khác, giữ lại từ đầu tiên của key, thay bằng value
##    chuyển đoạn văn từ list về đoạn
##    
##    input: key, từ để đổi, danh sách vị trí đổi, số thứ tự key, đoạn văn chứa key
##    output: đoạn văn đã được đổi từ ở vị trí chỉ định, số thứ tự key
    line_split = p.text.split() # split đoạn
    key_split = key.split() # split key
    len_key = len(key_split)
    for i in range(len(line_split)):
        if re.search(key_split[0],line_split[i],re.IGNORECASE):# nếu từ đầu trong key xuất hiện
            count = 0 # đếm từ trong key
            while count < len_key:
                if re.search(key_split[count],line_split[i+count],re.IGNORECASE): ##so khớp không phân biệt hoa thường
                    count+=1 # đếm xem các từ trong key xuất hiện đủ chưa
                else:
                    break
            if count == len_key: # nếu đủ
                countKey += 1
                #punctuation =""
                #if re.match(r'\S', p.text): #so khớp với ký tự không phải chữ
                    #punctuation = line_split[i+count-1][-1] # dấu câu
                if countKey in numberList:  # bắt đầu thay đổi ở các vị trí cần thiết
                    count_1 = 1
                    while count_1 < len_key:
                        line_split[i+count_1] = "" #thêm u ở phía trước để xử lý ký tự tiếng việt nhá
                        count_1+=1
                    line_split[i] = value #+punctuation
                p.text = u" ".join(line_split)
                p.text = u" ".join(p.text.split()) #loai bỏ khoảng trắng trùng lặp
    return countKey

def replace(filename,key,value,numberList,output_file):
##    hàm duyệt từng đoạn trong file
##    tìm và thay thế từ ở vị trí chỉ định
##    input: tên file, từ muốn đổi, từ để đổi, danh sách vị trí đổi
##    output: file word đã được thay từ ở những vị trí chỉ định
    countKey = 0 # khởi tạo số thứ tự key
    doc = Document(filename)
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if re.search(key,block.text,re.IGNORECASE): #so khớp không phân biệt hoa thường
                countKey = replace_string(key,value,numberList,countKey,block)
        else:
            for row in block.rows:
                for p in row.cells:
                    if re.search(key,p.text,re.IGNORECASE): #so khớp không phân biệt hoa thường
                        countKey = replace_string(key,value,numberList,countKey,p)    
    doc.save(output_file)

'''input_file = 'output/phong8.docx'
input_file = os.path.abspath(input_file)
#file = os.getcwd() + "/" + input_file
key = u'công việc'
value = u'công việc tuần'
numberList=[1,2]
output_file = 'output/output.docx'
replace(input_file,key,value,numberList,output_file)'''